import json
import unicodedata
import base64
from http.server import BaseHTTPRequestHandler
from io import BytesIO
from docx import Document

def remover_acentos(texto):
    if not texto: return ""
    nfkd = unicodedata.normalize('NFKD', texto)
    return "".join([c for c in nfkd if not unicodedata.combining(c)])

def extrair_dados_do_docx(conteudo_arquivo):
    try:
        doc = Document(BytesIO(conteudo_arquivo))
        texto_bruto = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt: texto_bruto.append(txt)
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    txt = celula.text.strip()
                    if txt and txt not in texto_bruto: texto_bruto.append(txt)
        
        lista_final = []
        for i in range(0, len(texto_bruto), 2):
            if i + 1 < len(texto_bruto):
                pergunta = texto_bruto[i]
                resposta = remover_acentos(texto_bruto[i+1].upper().replace(" ", ""))
                lista_final.append({"pergunta": pergunta, "resposta": resposta})
        return lista_final
    except Exception as e:
        print(f"Erro no processamento do Docx: {str(e)}")
        return []

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers.get('Content-Length', 0))
            post_data = self.rfile.read(content_length)
            
            # Carrega o JSON enviado pelo Next.js
            dados_json = json.loads(post_data.decode('utf-8'))
            base64_string = dados_json.get('arquivoBase64', '')
            
            # Converte o texto Base64 de volta para os bytes reais do arquivo Word
            bytes_arquivo = base64.b64decode(base64_string)
            
            # Extrai as perguntas
            questoes = extrair_dados_do_docx(bytes_arquivo)
            
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            
            response = {"success": True, "questoes": questoes}
            self.wfile.write(json.dumps(response).encode('utf-8'))
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"success": False, "error": str(e)}).encode('utf-8'))
        return

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
        return
