import streamlit as st
import unicodedata
import random
import time
import os
import string
from io import BytesIO
from docx import Document
from supabase import create_client

# ==================================================
# 1. CONEXÃO E CONFIGURAÇÃO
# ==================================================
URL_SUPABASE = st.secrets["URL_SUPABASE"]
KEY_SUPABASE = st.secrets["KEY_SUPABASE"]
supabase = create_client(URL_SUPABASE, KEY_SUPABASE)

st.set_page_config(page_title="Arena da Forca", page_icon="⚔️", layout="wide")

# --- FUNÇÕES DE APOIO ---
def remover_acentos(texto):
    if not texto: 
        return ""
    nfkd = unicodedata.normalize('NFKD', texto)
    return "".join([c for c in nfkd if not unicodedata.combining(c)])

def extrair_dados_do_docx(arquivo_docx):
    try:
        doc = Document(arquivo_docx)
        texto_bruto = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt: 
                texto_bruto.append(txt)
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    txt = celula.text.strip()
                    if txt and txt not in texto_bruto: 
                        texto_bruto.append(txt)
        lista_final = []
        for i in range(0, len(texto_bruto), 2):
            if i + 1 < len(texto_bruto):
                pergunta = texto_bruto[i]
                resposta = remover_acentos(texto_bruto[i+1].upper().replace(" ", ""))
                lista_final.append({"pergunta": pergunta, "resposta": resposta})
        return lista_final
    except Exception as e:
        st.error(f"Erro ao ler o documento Word: {e}")
        return []

def gerar_senha_aleatoria():
    caracteres = string.ascii_uppercase + string.digits
    return "".join(random.choice(caracteres) for _ in range(4))

# ==================================================
# 2. LOGIN E ESTADO DO JOGO
# ==================================================
if "jogador" not in st.session_state:
    st.session_state.jogador = None

if "clique_bloqueado" not in st.session_state:
    st.session_state.clique_bloqueado = False

if not st.session_state.jogador:
    st.title("⚔️ Arena da Forca")
    perfil = st.radio("Escolha seu perfil para entrar:", ["Jogador", "Mestre do Jogo (Admin)"])
    
    if perfil == "Mestre do Jogo (Admin)":
        st.info("Acesso restrito. Insira as credenciais do Mestre para assumir o controle.")
        senha_admin = st.text_input("Digite a chave de acesso do Mestre:", type="password", key="input_senha_admin")
        
        if st.button("ENTRAR COMO MESTRE"):
            if senha_admin and senha_admin.strip().upper() == "TREINAMENTOWLI":
                st.session_state.jogador = "TREINAMENTOWLI"
                nova_senha = gerar_senha_aleatoria()
                try:
                    supabase.table("forca_disputa_arena").update({"forca_senha_acesso": nova_senha}).eq("id", 1).execute()
                except Exception:
                    pass
                st.rerun()
            else:
                st.error("🔒 Chave de acesso do Mestre incorreta ou negada.")
                
    else:
        nome = st.text_input("Digite seu nome para entrar na Arena:", key="input_nome")
        senha_digitada = st.text_input("Digite a Senha da Arena:", type="password", key="input_senha")
        
        if st.button("ENTRAR NA ARENA"):
            if nome and senha_digitada:
                nome_upper = nome.strip().upper()
                
                if nome_upper == "TREINAMENTOWLI":
                    st.error("Use a opção de perfil Administrador para acessar com este nome.")
                else:
                    try:
                        res_arena = supabase.table("forca_disputa_arena").select("forca_senha_acesso").eq("id", 1).single().execute()
                        senha_valida = res_arena.data.get('forca_senha_acesso', '1234') if res_arena.data else '1234'
                    except Exception:
                        senha_valida = '1234'

                    if senha_digitada.strip().upper() == senha_valida.upper():
                        st.session_state.jogador = nome_upper
                        check_user = supabase.table("forca_disputa_ranking").select("*").eq("jogador", nome_upper).execute()
                        
                        if not check_user.data:
                            total_users = supabase.table("forca_disputa_ranking").select("jogador", count="exact").neq("jogador", "TREINAMENTOWLI").execute()
                            qtd_atual = total_users.count if total_users.count is not None else 0
                            supabase.table("forca_disputa_ranking").upsert(
                                {"jogador": nome_upper, "pontos": 0, "forca_avatar_num": qtd_atual + 1}, 
                                on_conflict="jogador"
                            ).execute()
                        st.rerun()
                    else:
                        st.error("🔑 Senha incorreta.")
            else:
                st.warning("Preencha todos os campos obrigatórios.")
    st.stop()

# ==================================================
# 3. LÓGICA DE JOGO CORE
# ==================================================
def calcular_proximo_turno(jogador_atual):
    try:
        res = supabase.table("forca_disputa_ranking").select("jogador").neq("jogador", "TREINAMENTOWLI").order("jogador").execute()
        lista_jogadores = [r['jogador'] for r in res.data]
        if not lista_jogadores: 
            return ""
        if len(lista_jogadores) == 1: 
            return lista_jogadores
        if jogador_atual in lista_jogadores:
            idx = lista_jogadores.index(jogador_atual)
            return lista_jogadores[(idx + 1) % len(lista_jogadores)]
        return lista_jogadores
    except Exception:
        return ""

def registrar_jogada(letra, jogo_atual):
    if st.session_state.jogador == "TREINAMENTOWLI": 
        return
    st.session_state.clique_bloqueado = True
    lista_antiga = jogo_atual.get('letras_tentadas', '')
    tentadas = [l.strip() for l in lista_antiga.split(",") if l.strip()]
    
    if letra in tentadas:
        st.session_state.clique_bloqueado = False
        return

    novas_letras = (lista_antiga + "," + letra) if lista_antiga else letra
    novos_erros = jogo_atual.get('erros', 0)
    palavra_alvo = jogo_atual.get('palavra', 'ARENA')
    modo_jogo = jogo_atual.get('forca_modo_jogo', 'LIVRE')
    
    res_p = supabase.table("forca_disputa_ranking").select("pontos").eq("jogador", st.session_state.jogador).single().execute()
    pts_atuais = res_p.data["pontos"] if res_p.data else 0

    if letra in palavra_alvo:
        supabase.table("forca_disputa_ranking").update({"pontos": pts_atuais + 5}).eq("jogador", st.session_state.jogador).execute()
    else:
        novos_erros += 1
        if modo_jogo == "TURNOS":
            supabase.table("forca_disputa_ranking").update({"pontos": max(0, pts_atuais - 5)}).eq("jogador", st.session_state.jogador).execute()
    
    proximo = calcular_proximo_turno(st.session_state.jogador)
    supabase.table("forca_disputa_arena").update({
        "letras_tentadas": novas_letras,
        "erros": novos_erros,
        "ultimo_jogador": st.session_state.jogador,
        "forca_proximo_turno": proximo,
        "forca_timestamp_inicio": time.time()
    }).eq("id", 1).execute()

def reiniciar_arena_completa():
    try:
        supabase.table("forca_disputa_ranking").update({"pontos": 0}).neq("jogador", "TREINAMENTOWLI").execute()
        res = supabase.table("forca_disputa_arena").select("forca_modo_jogo", "forca_senha_acesso", "forca_tempo_maximo").eq("id", 1).single().execute()
        dados = res.data if res.data else {"forca_modo_jogo": "LIVRE", "forca_senha_acesso": "1234", "forca_tempo_maximo": 15}
        
        supabase.table("forca_disputa_arena").update({
            "pergunta": "Aguardando nova pergunta...", "palavra": "ARENA",
            "letras_tentadas": "", "erros": 0, "restantes": 0, "ultimo_jogador": "SISTEMA",
            "forca_modo_jogo": dados.get("forca_modo_jogo"), "forca_senha_acesso": dados.get("forca_senha_acesso"),
            "forca_proximo_turno": "", "forca_tempo_maximo": dados.get("forca_tempo_maximo"), "forca_timestamp_inicio": 0.0
        }).eq("id", 1).execute()
    except Exception:
        pass
    st.session_state.clique_bloqueado = False
    st.rerun()
# ==================================================
# 4. INTERFACE GRÁFICA VIVA DA ARENA
# ==================================================
@st.fragment(run_every=1)
def arena_viva():
    st.session_state.clique_bloqueado = False
    if "podio_liberado" not in st.session_state: st.session_state.podio_liberado = False
    if "rodada_terminada" not in st.session_state: st.session_state.rodada_terminada = False

    try:
        # Buscamos o registro e guardamos o resultado bruto
        resposta_banco = supabase.table("forca_disputa_arena").select("*").eq("id", 1).single().execute()
        jogo = resposta_banco.data
    except Exception as erro_conexao:
        # CORREÇÃO: Mostra o erro real na tela caso falhe
        st.warning(f"Erro de Conexão com a Arena: {erro_conexao}")
        return
        
    if not jogo:
        st.warning("A linha com ID = 1 não foi encontrada na tabela forca_disputa_arena.")
        return


    # Execução controlada de áudio para evitar bugs de loop
    if jogo['pergunta'] != "Aguardando nova pergunta..." and jogo['erros'] < 6:
        if os.path.exists("musica.mp3") and "tocando_musica" not in st.session_state:
            st.audio("musica.mp3", format="audio/mp3", loop=True, autoplay=True)
            st.session_state.tocando_musica = True

    c_img, c_txt = st.columns([1, 4])
    erros_atuais = jogo.get('erros', 0)
    ultimo_player = jogo.get('ultimo_jogador', "SISTEMA")
    modo_jogo = jogo.get('forca_modo_jogo', "LIVRE")
    proximo_autorizado = jogo.get('forca_proximo_turno', "")
    tempo_maximo = jogo.get('forca_tempo_maximo', 15)
    timestamp_inicio = jogo.get('forca_timestamp_inicio', 0.0)
    
    with c_img:
        nome_img = f"erro{erros_atuais}.png"
        if os.path.exists(nome_img): 
            st.image(nome_img, width=140)
        else: 
            st.metric("Erros", f"{erros_atuais}/6")

    with c_txt:
        contagem = jogo.get('restantes', 0)
        tentadas = [l.strip() for l in jogo['letras_tentadas'].split(",") if l.strip()]
        palavra_alvo = jogo['palavra']
        vitoria = all((letra == " " or letra in tentadas) for letra in palavra_alvo)

        st.session_state.rodada_terminada = bool(vitoria or erros_atuais >= 6)

        if vitoria and erros_atuais < 6:
            id_v = f"vitoria_{palavra_alvo}_{contagem}"
            if id_v not in st.session_state:
                if st.session_state.jogador == ultimo_player and st.session_state.jogador != "TREINAMENTOWLI":
                    try:
                        res_p = supabase.table("forca_disputa_ranking").select("pontos").eq("jogador", st.session_state.jogador).single().execute()
                        pts = res_p.data["pontos"] if res_p.data else 0
                        supabase.table("forca_disputa_ranking").update({"pontos": pts + 10}).eq("jogador", st.session_state.jogador).execute()
                    except Exception: 
                        pass
                st.session_state[id_v] = True
                
        if (vitoria or erros_atuais >= 6) and contagem == 0:
            st.error("💀 DESAFIO ENCERRADO!")
        else:
            prefixo = f"📝 Pergunta {contagem}" if contagem > 0 else "🔥 DESAFIO FINAL"
            st.markdown(f"### {prefixo}")
            st.info(f"❓ {jogo['pergunta']}")

        texto_visual = "".join([f"{l} " if (l == " " or l in tentadas or erros_atuais >= 6) else "_ " for l in palavra_alvo])
        st.markdown(f"```\n{texto_visual}\n```")
        st.caption(f"Último lance por: {ultimo_player} | Modo: {modo_jogo}")

    autorizado_a_jogar = True
    if modo_jogo == "TURNOS" and not vitoria and erros_atuais < 6:
        if proximo_autorizado:
            segundos_restantes = max(0, int(tempo_maximo - (time.time() - timestamp_inicio))) if timestamp_inicio else tempo_maximo
            
            if segundos_restantes <= 0:
                try:
                    res_punido = supabase.table("forca_disputa_ranking").select("pontos").eq("jogador", proximo_autorizado).single().execute()
                    if res_punido.data:
                        supabase.table("forca_disputa_ranking").update({"pontos": max(0, res_punido.data["pontos"] - 5)}).eq("jogador", proximo_autorizado).execute()
                except Exception: 
                    pass
                supabase.table("forca_disputa_arena").update({
                    "forca_proximo_turno": calcular_proximo_turno(proximo_autorizado), 
                    "forca_timestamp_inicio": time.time(),
                    "ultimo_jogador": f"SISTEMA (TEMPO DE {proximo_autorizado} ESGOTOU)"
                }).eq("id", 1).execute()
            else:
                st.progress(segundos_restantes / tempo_maximo, text=f"⏱️ Tempo restante: {segundos_restantes}s")
                if st.session_state.jogador == proximo_autorizado:
                    st.success("⚔️ É a sua vez de jogar!")
                else:
                    st.warning(f"⏳ Aguarde sua vez. Jogador atual: {proximo_autorizado}")
                    autorizado_a_jogar = False

    if not vitoria and erros_atuais < 6:
        letras_abc = "ABCDEFGHIJKLMNOPQRSTUVWXYZ-"
        cols_tec = st.columns(13)
        for i, letra in enumerate(letras_abc):
            ja_foi = letra in tentadas
            # O .get() retorna None se a variável sumir, impedindo o app de quebrar
            is_admin = st.session_state.get("jogador") == "TREINAMENTOWLI"

            disabled = ja_foi or (not autorizado_a_jogar) or st.session_state.clique_bloqueado or is_admin
            
            if cols_tec[i % 13].button(letra, key=f"btn_letra_{letra}", disabled=disabled, use_container_width=True):
                st.session_state.clique_bloqueado = True
                if modo_jogo == "TURNOS" and not proximo_autorizado:
                    supabase.table("forca_disputa_arena").update({"forca_timestamp_inicio": time.time()}).eq("id", 1).execute()
                registrar_jogada(letra, jogo)

    if st.session_state.jogador != "TREINAMENTOWLI":
        st.divider()
        st.markdown("### 🏆 Placar Global")
        try:
            res_rank = supabase.table("forca_disputa_ranking").select("*").order("pontos", desc=True).execute().data
            j_competidores = [r for r in res_rank if r['jogador'] != "TREINAMENTOWLI"]
            if j_competidores:
                cols_r_p = st.columns(min(len(j_competidores), 5))
                for idx, r in enumerate(j_competidores[:10]):
                    with cols_r_p[idx % 5]:
                        st.write(f"**{idx+1}º {r['jogador']}**")
                        avatar = f"AV{r.get('forca_avatar_num')}.png"
                        if os.path.exists(avatar): 
                            st.image(avatar, width=40)
                        else: 
                            st.write("👤")
                        st.caption(f"{r['pontos']} pts")
        except Exception:
            pass

# Executa a tela imediatamente caso seja um participante comum
if st.session_state.jogador and st.session_state.jogador != "TREINAMENTOWLI":
    st.title("⚔️ Arena da Forca")
    arena_viva()
# ==================================================
# 5. PAINEL ADM (FORA DO CONTEXTO DE FRAGMENTO)
# ==================================================
if st.session_state.jogador == "TREINAMENTOWLI":
    st.title("⚔️ Painel do Mestre - Arena da Forca")
    
    # --- SINCRONIZAÇÃO GERAL DE ESTADOS INDEPENDENTE ---
    try:
        jogo_check = supabase.table("forca_disputa_arena").select("*").eq("id", 1).single().execute().data
        senha_atual = jogo_check.get('forca_senha_acesso', '----') if jogo_check else '----'
        
        if jogo_check:
            tentadas_check = [l.strip() for l in jogo_check.get('letras_tentadas', '').split(",") if l.strip()]
            palavra_check = jogo_check.get('palavra', 'ARENA')
            erros_check = jogo_check.get('erros', 0)
            contagem_restante = jogo_check.get('restantes', 0)
            
            vitoria_check = all((letra == " " or letra in tentadas_check) for letra in palavra_check)
            
            # O pódio é validado se o desafio encerrou E não há mais questões na fila
            if (vitoria_check or erros_check >= 6) and contagem_restante == 0:
                st.session_state.rodada_terminada = True
            else:
                st.session_state.rodada_terminada = False
    except Exception:
        jogo_check, senha_atual = None, '----'
        
    url_completa = "https://streamlit.io"

    if "podio_liberado" not in st.session_state: 
        st.session_state.podio_liberado = False

    def avancar_proxima_pergunta():
        if "fila_perguntas" in st.session_state and st.session_state.fila_perguntas:
            proxima = st.session_state.fila_perguntas.pop(0)
            try:
                supabase.table("forca_disputa_arena").update({
                    "pergunta": proxima['pergunta'], 
                    "palavra": proxima['resposta'],
                    "letras_tentadas": "", 
                    "erros": 0, 
                    "restantes": len(st.session_state.fila_perguntas),
                    "ultimo_jogador": "SISTEMA", 
                    "forca_proximo_turno": "", 
                    "forca_timestamp_inicio": 0.0
                }).eq("id", 1).execute()
            except Exception:
                pass
            st.session_state.podio_liberado = False
            st.session_state.rodada_terminada = False
            st.rerun()

    # Criação das Abas
    abas = st.tabs(["🎮 ARENA DO JOGO", "👥 CONTROLE DE PARTICIPANTES", "📱 QR CODE", "🏆 PODER DOS CAMPEÕES"])

    # Força o clique automático na quarta aba do navegador assim que liberado
    if st.session_state.get('rodada_terminada', False) and st.session_state.get('podio_liberado', False):
        st.components.v1.html(
            "<script>window.parent.document.querySelectorAll('button[role=\"tab\"]').click();</script>", 
            height=0
        )

    # --- ABA 0: CONTEÚDO EXCLUSIVO DA ARENA DO JOGO ---
    with abas:
        col_tab, col_menu = st.columns(2)
        with col_tab:
            arena_viva()
        with col_menu:
            if st.button("➡️ Próxima", use_container_width=True, key="btn_prox_mestre"):
                avancar_proxima_pergunta()
            st.divider()
            st.markdown("### 🏆 Ranking")
            try:
                res_rank = supabase.table("forca_disputa_ranking").select("*").execute().data
                if res_rank:
                    res_rank_ord = sorted(res_rank, key=lambda x: x.get('pontos', x.get('points', 0)), reverse=True)
                    jogadores_f = [r for r in res_rank_ord if r.get('jogador') != "TREINAMENTOWLI"]
                    for i, r in enumerate(jogadores_f[:10]):
                        pts = r.get('pontos', r.get('points', 0))
                        st.write(f"{i+1}º {r['jogador']}: **{pts} pts**")
                else: 
                    st.write("Nenhum competidor na arena.")
            except Exception: 
                st.write("Sincronizando placar...")

        st.divider()
        with st.expander("⚙️ CONFIGURAÇÃO DE QUESTÕES", expanded=True):
            if "fila_perguntas" not in st.session_state: 
                st.session_state.fila_perguntas = []
            c1, c2 = st.columns(2)
            with c1:
                arquivo = st.file_uploader("Arquivo .docx", type=["docx"], key="uploader_doc")
                if st.button("📥 PROCESSAR ARQUIVO"):
                    if arquivo:
                        st.session_state.fila_perguntas = extrair_dados_do_docx(arquivo)
                        st.success(f"{len(st.session_state.fila_perguntas)} questões carregadas!")
            with c2:
                st.metric("Fila Restante", len(st.session_state.fila_perguntas))
                if jogo_check:
                    novo_t = st.number_input("Tempo limite (s):", min_value=5, max_value=120, value=int(jogo_check.get('forca_tempo_maximo', 15)))
                    if novo_t != jogo_check.get('forca_tempo_maximo'):
                        supabase.table("forca_disputa_arena").update({"forca_tempo_maximo": novo_t}).eq("id", 1).execute()
                        st.rerun()
                
                st.write("")
                modo_banco = jogo_check.get('forca_modo_jogo', 'LIVRE') if jogo_check else 'LIVRE'
                index_modo = 0 if modo_banco == "LIVRE" else 1
                novo_modo = st.radio("Alternar Formato de Jogo:", ["LIVRE", "TURNOS"], index=index_modo)
                
                if novo_modo != modo_banco:
                    supabase.table("forca_disputa_arena").update({"forca_modo_jogo": novo_modo, "forca_proximo_turno": ""}).eq("id", 1).execute()
                    st.rerun()
                
                st.write("")
                if st.button("🔄 REINICIAR ARENA COMPLETA", use_container_width=True):
                    reiniciar_arena_completa()
    # --- ABA 1: GERENCIAMENTO DE PARTICIPANTES ---
    with abas:
        st.markdown("### 👥 Gerenciamento de Participantes na Sala")
        if st.button("🗑️ EXPULSAR TODOS OS JOGADORES DA ARENA", use_container_width=True, type="primary"):
            supabase.table("forca_disputa_ranking").delete().neq("jogador", "TREINAMENTOWLI").execute()
            supabase.table("forca_disputa_arena").update({"forca_proximo_turno": ""}).eq("id", 1).execute()
            st.rerun()
            
        res_j = supabase.table("forca_disputa_ranking").select("jogador").neq("jogador", "TREINAMENTOWLI").execute().data
        if res_j:
            for j in res_j:
                c1, c2 = st.columns(2)
                c1.markdown(f"__👤 {j['jogador']}__")
                if c2.button("❌ EXPULSAR", key=f"excluir_{j['jogador']}", use_container_width=True):
                    supabase.table("forca_disputa_ranking").delete().eq("jogador", j['jogador']).execute()
                    st.rerun()

    # --- ABA 2: CONEXÃO VIA QR CODE ---
    with abas:
        st.markdown(f"<h1 style='text-align:center; color:#3b82f6; font-family:monospace;'>Chave: {senha_atual}</h1>", unsafe_allow_html=True)
        col_esq_qr, col_cen_qr, col_dir_qr = st.columns(3)
        with col_cen_qr:
            if os.path.exists("QRCode Forca.png"):
                st.image("QRCode Forca.png", use_container_width=True)
            else:
                st.error("⚠️ O arquivo 'QRCode Forca.png' não foi localizado no diretório atual.")

    # --- ABA 3: PODER DOS CAMPEÕES (PÓDIO ESCALONADO AUTOMÁTICO) ---
    with abas:
        if st.session_state.get('rodada_terminada', False) and not st.session_state.podio_liberado:
            if st.button("🏆 LIBERAR EXIBIÇÃO DOS CAMPEÕES NO TELÃO", type="primary", use_container_width=True, key="btn_mestre_liberar_podio"):
                st.session_state.podio_liberado = True
                st.rerun()
                
        elif st.session_state.get('rodada_terminada', False) and st.session_state.podio_liberado:
            st.markdown("<h1 style='text-align: center; color: #ffb703; font-size: 42px; margin-bottom: 20px;'>🏆 PÓDIO DA ARENA DA FORCA 🏆</h1>", unsafe_allow_html=True)
            st.write("")
            
            try:
                res_v = supabase.table("forca_disputa_ranking").select("*").neq("jogador", "TREINAMENTOWLI").execute().data
                if res_v:
                    # Ordena os jogadores por pontos (maior para o menor)
                    ranking_completo = sorted(res_v, key=lambda x: x.get('pontos', x.get('points', 0)), reverse=True)
                    total_jogadores = len(ranking_completo)
                    
                    # Cria a quantidade exata de colunas necessárias baseada nos jogadores reais
                    qtd_colunas = 3 if total_jogadores >= 3 else total_jogadores
                    cols_podio = st.columns(qtd_colunas)
                    
                    # FÓRMULA ANTIFILTRO: Calcula as posições sem usar listas estáticas de colchetes
                    for idx_coluna in range(qtd_colunas):
                        # Se houver 3 jogadores, organiza: 2º Lugar (coluna 0), 1º Lugar (coluna 1), 3º Lugar (coluna 2)
                        if qtd_colunas == 3:
                            if idx_coluna == 0: idx_ranking = 1
                            elif idx_coluna == 1: idx_ranking = 0
                            else: idx_ranking = 2
                        else:
                            # Se houver menos de 3, exibe de forma sequencial linear normal
                            idx_ranking = idx_coluna
                            
                        if idx_ranking < total_jogadores:
                            jogador_dados = ranking_completo[idx_ranking]
                            nome_jogador = jogador_dados['jogador']
                            pts_jogador = jogador_dados.get('pontos', jogador_dados.get('points', 0))
                            avatar_num = jogador_dados.get("forca_avatar_num", None)
                            arquivo_av = f"AV{avatar_num}.png" if avatar_num else None
                            
                            # Escalonamento rígido proporcional solicitado: 1º > 2º > 3º
                            if idx_ranking == 0:
                                label_colocacao = "👑 1º LUGAR"
                                cor_texto = "#10b981"
                                tamanho_avatar = 280
                            elif idx_ranking == 1:
                                label_colocacao = "🥈 2º LUGAR"
                                cor_texto = "#3b82f6"
                                tamanho_avatar = 200
                            else:
                                label_colocacao = "🥉 3º LUGAR"
                                cor_texto = "#64748b"
                                tamanho_avatar = 140
                                
                            with cols_podio[idx_coluna]:
                                st.markdown(f"<h3 style='text-align: center; color: {cor_texto};'>{label_colocacao}</h3>", unsafe_allow_html=True)
                                
                                if arquivo_av and os.path.exists(arquivo_av):
                                    st.image(arquivo_av, width=tamanho_avatar)
                                else:
                                    st.markdown(f"<p style='text-align: center; font-size: {tamanho_avatar//3}px;'>👤</p>", unsafe_allow_html=True)
                                    
                                st.markdown(f"""
                                <div style="text-align: center; margin-top: 10px;">
                                    <h2 style="font-size: 24px; margin-bottom: 2px;">{nome_jogador}</h2>
                                    <p style="font-size: 16px; color: #64748b; font-family: monospace; font-weight: bold;">{pts_jogador} PTS</p>
                                </div>
                                """, unsafe_allow_html=True)
            except Exception:
                st.error("Erro ao gerar os dados visuais do pódio.")
        else:
            st.info("O Pódio dos Campeões será montado automaticamente aqui assim que a Arena da Forca for encerrada pelo Mestre.")
